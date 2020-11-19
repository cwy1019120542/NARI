import os
import time
import uuid
import smtplib
import zipfile
import openpyxl
import random
import re
import shutil
from urllib.parse import quote
from sqlalchemy import and_
import pandas
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.utils import formatdate
from email.header import Header
from flask import session, jsonify, send_from_directory, make_response
from functools import wraps
from .extention import db, redis, celery
from .models import User, MainConfig, SendConfig, ReceiveConfig
from .parameter_config import accept_file_type
import sys
sys.path.append('/usr/local/lib/python3.8/dist-packages/python_docx-0.8.10-py3.8.egg')
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches
from docx.oxml.ns import qn


def clean_file_name(file_name):
    sub_list = r"[/\\\:\*\?\"\<\>\|\[\]《》\（\）-]"
    new_file_name = re.sub(sub_list, '_', str(file_name))
    return new_file_name


def to_xlsx(file_path):
    file_name, suffix = os.path.splitext(file_path)
    excel_file = pandas.read_excel(file_path)
    excel_file.to_excel(f'{file_name}.xlsx', index=False)
    os.remove(file_path)

def response(is_succcess, status_code, message, result=[], **kwargs):
    response_json = {
        "is_success": is_succcess,
        "status_code": status_code,
        "message": message,
        "result": result
    }
    response_json.update(kwargs)
    response_obj = jsonify(response_json)
    response_obj.status_code = status_code
    response_obj.headers['Access-Control-Allow-Origin'] = '*'
    # response_obj.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE'
    response_obj.headers['Cache-Control'] = "no-cache"
    return response_obj

def parameter_check(request_data, parameter_group_list, is_all=True):
    if request_data == None:
        return False, response(False, 400, "请求体为空")
    clean_data = {}
    if is_all:
        judge_list = [i for i, j, k, l in parameter_group_list if i not in request_data]
        if judge_list:
            return False, response(False, 400, "参数缺失")
    for parameter, parameter_type, is_none, max_len in parameter_group_list:
        if parameter not in request_data and not is_all:
            continue
        value = request_data[parameter]
        if value == None:
            if not is_none:
                return False, response(False, 400, "提供非法None值")
            else:
                continue
        if not isinstance(value, parameter_type):
            if parameter_type == bool:
                try:
                    value = bool(int(value))
                except:
                    return False, response(False, 400, "数据格式错误")
            try:
                value = parameter_type(value)
            except:
                return False, response(False, 400, "数据格式错误")
        if len(str(value)) > max_len:
            return False, response(False, 400, "字段长度超过限制")
        clean_data[parameter] = value
    return True, clean_data

def is_login(func):
    @wraps(func)
    def wrapper(user_id=None, *args, **kwargs):
        if user_id != None and session.get('user_id') != user_id:
            return response(False, 403, f"用户{user_id}未登录")
        return func(user_id, *args, **kwargs)
    return wrapper

def page_filter(model, clean_data, fuzzy_field):
    limit = clean_data.pop('limit', 10)
    offset = clean_data.pop('offset', 0)
    page_info = {
        "limit": limit,
        "offset": offset,
    }
    filter_query_list = and_(*[getattr(model, i)==clean_data[i] if i not in fuzzy_field else getattr(model, i).contains(clean_data[i]) for i in clean_data.keys()])
    all_data_query = db.session.query(model).filter(filter_query_list)
    data_count = all_data_query.count()
    data_list = all_data_query.order_by(-model.change_timestamp, -model.id).limit(limit).offset(offset).all()
    page_info["count"] = data_count
    return data_list, page_info

def captcha_check(email, captcha):
    if not redis.exists(email):
        return "验证码过期"
    db_captcha = redis.get(email).decode()
    if int(db_captcha) != int(captcha):
        return "验证码错误"

def resource_limit(resource_group):
    if not resource_group:
        return False, response(False, 400, "资源组出错")
    father_id = None
    resource_query = None
    resource = None
    for resource_index, (model, resource_id, link_field) in enumerate(resource_group):
        if not resource_id:
            break
        resource_query = db.session.query(model).filter_by(id=resource_id)
        resource = resource_query.first()
        if not resource:
            return False, response(False, 404, "请求的资源不存在")
        if resource_index > 0:
            if getattr(resource, link_field) != father_id:
                return False, response(False, 403, f"没有权限访问该资源{resource_id}")
        if resource_index != len(resource_group) - 1:
            father_id = resource_id
    return True, (model, resource_query, resource, resource_id, link_field, father_id)

def save_file(request_parameter, request_file, is_reset, file_dir, new_file_name=None, file_type="excel"):
    if request_parameter not in request_file:
        return False, response(False, 400, "参数错误")
    if os.path.exists(file_dir):
        if is_reset:
            shutil.rmtree(file_dir)
            os.makedirs(file_dir)
    else:
        os.makedirs(file_dir)
    file = request_file.get(request_parameter)
    file_name = clean_file_name(file.filename.strip('"'))
    if not file_name:
        return False, response(False, 400, "参数错误")
    file_prefix, file_suffix = os.path.splitext(file_name)
    if file_type == "excel":
        if not file_name.endswith(accept_file_type):
            return False, response(False, 400, "文件格式错误")
        return_file_name = f"{file_prefix}.xlsx"
        if new_file_name:
            file_name = new_file_name + file_suffix
            file_prefix = new_file_name
            return_file_name = f"{new_file_name}.xlsx"
        file_path = os.path.join(file_dir, file_name)
        file.save(file_path)
        if not file_name.endswith(('.xlsx', '.XLSX')):
            to_xlsx(file_path)
        if file_name.endswith('.XLSX'):
            change_file_name = f"{file_prefix}.xlsx"
            os.rename(file_path, os.path.join(file_dir, change_file_name))
    elif file_type == "*":
        return_file_name = file_name
        if new_file_name:
            file_name = new_file_name + file_suffix
            return_file_name = file_name
        file.save(os.path.join(file_dir, file_name))
    return True, response(True, 200, "成功", return_file_name)

def file_resource(resource_group, file_dir, request_method, request_parameter, request_file, file_type="excel", file_path=None, is_reset=True, new_file_name=None):
    is_success, return_data = resource_limit(resource_group)
    if not is_success:
        return return_data
    is_exists = os.path.exists(file_dir)
    if request_method == "GET":
        if not os.path.exists(file_dir) or not os.listdir(file_dir):
            return response(False, 404, "请求的资源不存在")
        file_path = os.path.join(file_dir, os.listdir(file_dir)[0])
        return return_file(file_path)
    elif request_method == 'POST' or request_method == 'PUT':
        return save_file(request_parameter, request_file, is_reset, file_dir, new_file_name, file_type)[1]
    elif request_method == 'DELETE':
        if file_path:
            if not os.path.exists(file_path):
                return response(False, 404, "资源不存在")
            os.remove(file_path)
            return response(True, 204, "成功")
        else:
            if is_exists:
                file_list = os.listdir(file_dir)
                if file_list:
                    for file in file_list:
                        os.remove(os.path.join(file_dir, file))
                    return response(True, 204, "成功")
                else:
                    return response(False, 404, "请求的资源不存在")
            else:
                return response(False, 404, "请求的资源不存在")

def resource_manage(resource_group, request_method, request_args, request_json, parameter):
    is_success, return_data = resource_limit(resource_group)
    if not is_success:
        return return_data
    model, resource_query, resource, resource_id, link_field, father_id = return_data
    if request_method == 'GET':
        if resource_id:
            result = resource.get_info()
            return response(True, 200, "成功", result)
        config_parameter_get = parameter["GET"]
        fuzzy_field = parameter["fuzzy_field"]
        is_right, clean_data = parameter_check(request_args, config_parameter_get, False)
        if not is_right:
            return response(False, 400, clean_data)
        if father_id:
            clean_data[link_field] = father_id
        resource_list, page_info = page_filter(model, clean_data, fuzzy_field)
        result = [i.get_info() for i in resource_list]
        return response(True, 200, "成功", result, **page_info)
    elif request_method == 'POST':
        config_parameter_post = parameter["POST"]
        is_right, clean_data = parameter_check(request_json, config_parameter_post)
        if not is_right:
            return clean_data
        if father_id:
            clean_data[link_field] = father_id
        now_timestamp = int(time.time())
        clean_data['create_timestamp'] = now_timestamp
        clean_data['change_timestamp'] = now_timestamp
        new_resource = model(**clean_data)
        db.session.add(new_resource)
        db.session.commit()
        result = new_resource.get_info()
        return response(True, 200, "成功", result)
    elif request_method == 'PUT':
        config_parameter_post = parameter["POST"]
        is_right, clean_data = parameter_check(request_json, config_parameter_post, False)
        if not is_right:
            return clean_data
        clean_data['change_timestamp'] = int(time.time())
        resource_query.update(clean_data)
        db.session.commit()
        result = db.session.query(model).get(resource_id).get_info()
        return response(True, 201, "成功", result)
    elif request_method == 'DELETE':
        resource_query.delete()
        db.session.commit()
        return response(True, 204, "成功")


def dict_to_tuple(target_dict):
    target_list = list(target_dict.items())
    target_list.sort(key=lambda i: i[0])
    return tuple(target_list)

def create_task_id(task_name, **kwargs):
    task_id_group = []
    sort_kwargs = dict_to_tuple(kwargs)
    for key, value in sort_kwargs:
        task_id_group.append('%'.join([str(key), str(value)]))
    task_id = ''.join([task_name, '@', '#'.join(task_id_group), '@', str(uuid.uuid1())])
    return task_id

def split_task_id(task_id):
    task_name, kwargs_str, _ = task_id.split('@')
    kwargs = {i.split('%')[0]:i.split('%')[1] for i in kwargs_str.split('#') if i}
    return task_name, kwargs

def get_active_task():
    i = celery.control.inspect()
    active_tasks = i.active()
    active_task_dict = {}
    scheduled_tasks = i.scheduled()
    scheduled_task_dict = {}
    if active_tasks:
        task_id_list = [i['id'] for i in list(active_tasks.values())[0]]
        for task_id in task_id_list:
            task_name, kwargs = split_task_id(task_id)
            sort_kwargs = dict_to_tuple(kwargs)
            active_task_dict[(task_name, sort_kwargs)] = task_id
    if scheduled_tasks:
        task_id_list = [i['request']['id'] for i in list(scheduled_tasks.values())[0]]
        for task_id in task_id_list:
            task_name, kwargs = split_task_id(task_id)
            sort_kwargs = dict_to_tuple(kwargs)
            scheduled_task_dict[(task_name, sort_kwargs)] = task_id
    return active_task_dict, scheduled_task_dict

def get_run_config(user_id, target_key='main_config_id'):
    active_task, scheduled_task = get_active_task()
    active_task_key = list(active_task.keys())
    active_task_id_list = [int(i[target_key]) for i in [dict(j[1]) for j in active_task_key] if i.get('user_id') == str(user_id) and i.get(target_key)]
    scheduled_task_key = list(scheduled_task.keys())
    scheduled_task_id_list = [int(i[target_key]) for i in [dict(j[1]) for j in scheduled_task_key] if i.get('user_id') == str(user_id) and i.get(target_key)]
    return active_task_id_list, scheduled_task_id_list

def is_active_task(task_name, **kwargs):
    sort_kwargs = dict_to_tuple(kwargs)
    task_id_group = ((task_name, sort_kwargs))
    active_task_dict = get_active_task()
    if task_id_group in active_task_dict:
        return True

def get_file_name(config_files_dir, main_config_id, excel_dir_name, is_all=False):
    excel_dir = os.path.join(config_files_dir, str(main_config_id), excel_dir_name)
    if not os.path.exists(excel_dir):
        return None
    file_name_list = os.listdir(excel_dir)
    if not file_name_list:
        return None
    if is_all:
        return file_name_list
    else:
        return file_name_list[0]

def get_file_path(config_files_dir, main_config_id, excel_dir_name, is_all=False):
    excel_dir = os.path.join(config_files_dir, str(main_config_id), excel_dir_name)
    if not os.path.exists(excel_dir):
        return False, "所需表不存在"
    excel_list = os.listdir(excel_dir)
    if not excel_list:
        return False, "所需表不存在"
    if is_all:
        path_list = []
        for excel_name in excel_list:
            excel_path = os.path.join(excel_dir, excel_name)
            path_list.append(excel_path)
        return True, path_list
    excel_name = excel_list[0]
    excel_path = os.path.join(excel_dir, excel_name)
    return True, excel_path

def get_match_dict(config_files_dir, main_config_id):
    is_success, return_data = get_file_path(config_files_dir, main_config_id, "match_excel")
    if not is_success:
        return False, return_data
    match_excel = openpyxl.load_workbook(return_data, data_only=True)
    match_sheet = match_excel['Sheet1']
    total_data_list = list(match_sheet.values)
    match_dict = {}
    for data in total_data_list[1:]:
        target = data[0]
        email = data[1].strip() if data[1] else None
        if not email:
            continue
        if target in match_dict:
            if email not in match_dict[target]:
                match_dict[target].append(email)
        else:
            match_dict[target] = [email]
    match_excel.close()
    return True, match_dict

def check_dir(target_dir):
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)

def smtp_send_mail(smtp_obj, sender, receiver, subject, content, attachment_list=[]):
    msg = MIMEMultipart()
    msg['Subject'] = Header(subject, 'utf-8')
    msg['From'] = sender
    msg['To'] = receiver
    msg['Date'] = formatdate(localtime=True)
    text_plain = MIMEText(content, 'html', 'utf-8')
    msg.attach(text_plain)
    for attachment in attachment_list:
        sendfile = open(attachment, 'rb').read()
        text_att = MIMEText(sendfile, 'base64', 'utf-8')
        text_att["Content-Type"] = 'application/octet-stream'
        file_dir, file_name = os.path.split(attachment)
        # text_att["Content-Disposition"] = f'attachment; filename="{file_name}"'
        text_att.add_header('Content-Disposition', 'attachment', filename=clean_file_name(file_name))
        msg.attach(text_att)
    smtp_obj.sendmail(sender, receiver, msg.as_string())

def get_header_row(sheet_obj, header_row=None):
    if header_row:
        return int(header_row)
    data_list = sheet_obj.values
    for data_index, data in enumerate(data_list):
        if len([i for i in data if i]) > 2:
            return data_index + 1
    return 1

def get_column_number():
    letter_str = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    letter_list_single = list(letter_str)
    letter_list_double = [i+j for i in letter_str for j in letter_str]
    letter_list = letter_list_single + letter_list_double
    return letter_list

def return_file(file_path):
    base_dir, filename = os.path.split(file_path)
    response = make_response(send_from_directory(base_dir, filename))
    response.headers["Content-Disposition"] = "attachment; filename={0}; filename*=utf-8''{0}".format(quote(filename))
    return response

def send_multi_mail(ip, port, username, password, email_list, subject, content):
    try:
        smtp_obj = smtplib.SMTP(ip, port)
        smtp_obj.login(username, password)
    except Exception as error:
        print(error)
        return False, "运行失败,无法连接至发件邮箱"
    for email in email_list:
        try:
            smtp_send_mail(smtp_obj, username, email, subject, content)
        except Exception as error:
            print(error)
            continue
    # smtp_obj.quit()
    return True, "成功"

def get_task_info(task_id):
    if not task_id:
        return True, "未启动"
    task = celery.AsyncResult(task_id.decode())
    status = task.status
    result = task.result
    if status == 'SUCCESS':
        return True, result
    elif status == 'STARTED' or status == 'PENDING':
        return False, "正在运行"
    elif status == 'REVOKED':
        return True, "终止"

def return_zip(resource_group, zip_path, file_dir):
    is_success, return_data = resource_limit(resource_group)
    if not is_success:
        return return_data
    if not os.path.exists(file_dir) or not os.listdir(file_dir):
        return response(False, 404, file_dir)
    if os.path.exists(zip_path):
        os.remove(zip_path)
    zip_file = zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED)
    for file_name in os.listdir(file_dir):
        file_path = os.path.join(file_dir, file_name)
        zip_file.write(file_path, file_name)
    zip_file.close()
    return return_file(zip_path)

def start_task(app_config, main_config_info, function_name, task_function):
    main_config_id = main_config_info["id"]
    user_id = main_config_info["user_id"]
    config_info = main_config_info[f"{function_name}_config_info"]
    if config_info:
        config_id = config_info["id"]
        now_time = int(time.time())
        task_id = create_task_id(function_name, user_id=user_id, config_id=config_id,
                                      main_config_id=main_config_id)
        is_timing = config_info["is_timing"]
        start_time = config_info["start_timestamp"]
        countdown = 0
        if is_timing:
            if not start_time or start_time < now_time:
                return False, response(False, 400, "定时时间错误")
            countdown = start_time - int(time.time())
        redis.set(f'{main_config_id}_{function_name}_task_id', task_id)
        app_config_info = {
            "CONFIG_FILES_DIR": app_config["CONFIG_FILES_DIR"]
        }
        task_function.apply_async(kwargs={"app_config_info": app_config_info, "main_config_info": main_config_info},
                              task_id=task_id, countdown=countdown)
        return True, True
    else:
        return False, response(False, 400, "配置缺失")

def return_target_file(resource_group, file_path):
    if resource_group:
        is_success, return_data = resource_limit(resource_group)
        if not is_success:
            return return_data
    if not os.path.exists(file_path):
        return response(False, 404, "资源不存在")
    else:
        return return_file(file_path)

def handle_num(num):
    f_num = float(num)
    if f_num >= 10000 or f_num <= -10000:
        return str(round(num / 10000, 2)) + "亿"
    else:
        return str(round(f_num, 2)) + "万"

def get_report_data(workbook, sheet_name, sum_column_index, number, number_column_index, model=None):
    sheet = workbook[sheet_name]
    data_list = list(sheet.values)[1:]
    sum_data = data_list.pop()
    sum_list = [sum_data[i] for i in sum_column_index]
    sort_index = number_column_index[-1]
    data_list.sort(key=lambda x:float(str(x[sort_index] if x[sort_index] else 0).strip("%")), reverse=True)
    number_data_list = data_list[:number]
    part_data_list = [[i[j] for j in number_column_index] for i in number_data_list]
    if model == 1:
        model_amount = handle_num(sum(i[sort_index] for i in data_list if float(i[sort_index])>0))
        return sum_list, part_data_list, model_amount
    return sum_list, part_data_list

def connect_part_data(part_data_list, format_str="{}（{}元）", amount_index=1, company_index=0, name_dict=None):
    for part_data in part_data_list:
        if amount_index != None:
            amount_data = part_data[amount_index]
            part_data[amount_index] = handle_num(amount_data)
        company = part_data[company_index]
        if company in name_dict:
            part_data[company_index] = name_dict[company]
    return "、".join([format_str.format(*i) for i in part_data_list])

def judge_up_down(data, format_type=0):
    if format_type == 0:
        format_str = ("增加", "减少")
        handle_data = handle_num(data)
        compare_num = float(handle_data[:-1])
    else:
        format_str = ("增幅", "降幅")
        handle_data = float(data.strip("%") if data else 0)
        compare_num = handle_data
    format_data = str(handle_data).strip("-")
    if compare_num > 0:
        return f"{format_str[0]}{format_data}"
    else:
        return f"{format_str[1]}{format_data}"

def generate_rate(num1, num2):
    return round(float(num1) * 100 / float(num2), 2)

def generate_report_word(number, file_path, date, result_path):
    name_workbook = openpyxl.load_workbook("/home/cwy/static_files/公司全称与简称对照表.xlsx")
    name_sheet = name_workbook[name_workbook.sheetnames[0]]
    name_dict = dict(i[:2] for i in name_sheet.values)
    name_workbook.close()
    year, month = date.split("-")
    format_str_list = []
    format_str_list.append(None)
    format_str_list.append((month,))
    workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    sheet_name_list = workbook.sheetnames
    if len(sheet_name_list) < 18:
        return False, response(False, 400, "sheet页不匹配，请重新运行")
    sum_data_list1, part_data_list1, model_amount1 = get_report_data(workbook, "本月新增预开票情况统计表(全部)", [], number, [1, 4], model=1)
    sum_data_list2, part_data_list2, model_amount2 = get_report_data(workbook, "本月新增预开票情况统计表(系统内)", [], number, [1, 4], model=1)
    sum_data_list3, part_data_list3, model_amount3 = get_report_data(workbook, "本月新增滞后开票情况统计表(全部)", [], number, [1, 4], model=1)
    sum_data_list4, part_data_list4, model_amount4 = get_report_data(workbook, "本月新增滞后开票情况统计表(系统内)", [], number, [1, 4], model=1)

    format_str_list.append((month, model_amount1, model_amount2,
                            connect_part_data(part_data_list1, name_dict=name_dict), connect_part_data(part_data_list2, name_dict=name_dict),
                            model_amount3, model_amount4, connect_part_data(part_data_list3, name_dict=name_dict),
                            connect_part_data(part_data_list4, name_dict=name_dict)))
    format_str_list.append(None)
    sum_data_list5, part_data_list5 = get_report_data(workbook, "已开票未确认收入(预开票)余额清理情况统计表(全部)", [2, 4, 5, 6, 8, 9],
                                                      number, [1, 2])
    sum_data_list6, part_data_list6 = get_report_data(workbook, "已开票未确认收入(预开票)余额清理情况统计表(系统内)", [2, 4, 5, 6, 8, 9],
                                                      number, [1, 2])
    format_str_list.append((month, handle_num(sum_data_list5[0]), judge_up_down(sum_data_list5[1]),
                            judge_up_down(sum_data_list5[2], 1), handle_num(sum_data_list6[0]),
                            generate_rate(sum_data_list6[0], sum_data_list5[0]),
                            judge_up_down(sum_data_list6[1]),
                            judge_up_down(sum_data_list6[2], 1)))
    sum_data_list7, part_data_list7 = get_report_data(workbook, "已开票未确认收入(预开票)余额清理情况统计表(全部)", [],
                                                      number, [1, 6])
    sum_data_list8, part_data_list8 = get_report_data(workbook, "已开票未确认收入(预开票)余额清理情况统计表(系统内)", [],
                                                      number, [1, 6])
    format_str_list.append((connect_part_data(part_data_list5, name_dict=name_dict), handle_num(sum_data_list5[3]),
                            judge_up_down(sum_data_list5[4]), judge_up_down(sum_data_list5[5], 1),
                            connect_part_data(part_data_list7, name_dict=name_dict)))
    format_str_list.append((connect_part_data(part_data_list6, name_dict=name_dict), handle_num(sum_data_list6[3]),
                            judge_up_down(sum_data_list6[4]), judge_up_down(sum_data_list6[5], 1),
                            connect_part_data(part_data_list8, name_dict=name_dict)))
    format_str_list.append(None)
    sum_data_list9, part_data_list9 = get_report_data(workbook, "已开票未确认收入(滞后开票)余额清理情况统计表(全部)", [2, 4, 5, 6, 8, 9],
                                                      number, [1, 2])
    sum_data_list10, part_data_list10 = get_report_data(workbook, "已开票未确认收入(滞后开票)余额清理情况统计表(系统内)", [2, 4, 5, 6, 8, 9],
                                                        number, [1, 2])
    format_str_list.append((month, handle_num(sum_data_list9[0]), judge_up_down(sum_data_list9[1]),
                            judge_up_down(sum_data_list9[2], 1), handle_num(sum_data_list10[0]),
                            generate_rate(sum_data_list10[0], sum_data_list9[0]),
                            judge_up_down(sum_data_list10[1]),
                            judge_up_down(sum_data_list10[2], 1)))
    sum_data_list11, part_data_list11 = get_report_data(workbook, "已开票未确认收入(滞后开票)余额清理情况统计表(全部)", [],
                                                        number, [1, 6])
    sum_data_list12, part_data_list12 = get_report_data(workbook, "已开票未确认收入(滞后开票)余额清理情况统计表(系统内)", [],
                                                        number, [1, 6])
    format_str_list.append((connect_part_data(part_data_list9, name_dict=name_dict), handle_num(sum_data_list9[3]),
                            judge_up_down(sum_data_list9[4]), judge_up_down(sum_data_list9[5], 1),
                            connect_part_data(part_data_list11, name_dict=name_dict)))
    format_str_list.append((connect_part_data(part_data_list10, name_dict=name_dict), handle_num(sum_data_list10[3]),
                            judge_up_down(sum_data_list10[4]), judge_up_down(sum_data_list10[5], 1),
                            connect_part_data(part_data_list12, name_dict=name_dict)))
    format_str_list.extend([None, None, None, None])
    sum_data_list13, part_data_list13 = get_report_data(workbook, "项目成本结转不彻底", [2],
                                                        number, [1, 2])
    format_str_list.append((month, handle_num(sum_data_list13[0]), connect_part_data(part_data_list13, name_dict=name_dict)))
    format_str_list.extend([None, None])
    sum_data_list14, part_data_list14 = get_report_data(workbook, "生产成本长期挂账未结转情况统计表", [6, 8, 9],
                                                        number, [1, 6])
    sum_data_list15, part_data_list15 = get_report_data(workbook, "生产成本长期挂账未结转情况统计表", [],
                                                        number, [1, 6, 9])
    format_str_list.append((month, handle_num(sum_data_list14[0]), judge_up_down(sum_data_list14[1]),
                            judge_up_down(sum_data_list14[2], 1), connect_part_data(part_data_list14, name_dict=name_dict),
                            connect_part_data(part_data_list15, "{}（{}元, {}）", name_dict=name_dict)))
    format_str_list.extend([None, None, None])
    sum_data_list16, part_data_list16 = get_report_data(workbook, "挂账一年以上应付项目暂估情况分析", [2, 4, 5, 6, 8, 9],
                                                        number, [1, 2])
    sum_data_list17, part_data_list17 = get_report_data(workbook, "挂账一年以上应付项目暂估情况分析", [],
                                                        number, [1, 6])
    format_str_list.append((month, handle_num(sum_data_list16[0]), judge_up_down(sum_data_list16[1]),
                            judge_up_down(sum_data_list16[2], 1), connect_part_data(part_data_list16, name_dict=name_dict),
                            handle_num(sum_data_list16[3]), judge_up_down(sum_data_list16[4]),
                            judge_up_down(sum_data_list16[5], 1), connect_part_data(part_data_list17, name_dict=name_dict)))
    format_str_list.extend([None, None])
    sum_data_list18, part_data_list18 = get_report_data(workbook, "本月项目生产成本暂估比例异常(达20%以上)情况统计表", [3, 2, 6, 7],
                                                        number, [1, 2])
    sum_data_list19, part_data_list19 = get_report_data(workbook, "本月项目生产成本暂估比例异常(达20%以上)情况统计表", [],
                                                        number, [1, 3])
    format_str_list.append((
                           month, sum_data_list18[0], handle_num(sum_data_list18[1]), judge_up_down(sum_data_list18[2]),
                           judge_up_down(sum_data_list18[3], 1), connect_part_data(part_data_list18, name_dict=name_dict),
                           connect_part_data(part_data_list19, "{}（{}个）", None, name_dict=name_dict)))
    format_str_list.extend([None, None, None])
    sum_data_list20, part_data_list20 = get_report_data(workbook, "挂账一年以上应付原材料暂估情况分析", [2, 4, 5, 6, 8, 9],
                                                        number, [1, 6])
    format_str_list.append((month, handle_num(sum_data_list20[0]), judge_up_down(sum_data_list20[1]),
                            judge_up_down(sum_data_list20[2], 1), handle_num(sum_data_list20[3]),
                            judge_up_down(sum_data_list20[4]), judge_up_down(sum_data_list20[5], 1),
                            connect_part_data(part_data_list20, name_dict=name_dict)))
    format_str_list.append(None)
    sum_data_list21, part_data_list21 = get_report_data(workbook, "挂账一年以上预付账款情况", [6, 8, 9],
                                                        number, [1, 6])
    sum_data_list22, part_data_list22 = get_report_data(workbook, "挂账一年以上预付账款情况", [],
                                                        number, [1, 8, 9])
    format_str_list.append((month, handle_num(sum_data_list21[0]), judge_up_down(sum_data_list21[1]),
                            judge_up_down(sum_data_list21[2], 1), connect_part_data(part_data_list21, name_dict=name_dict),
                            connect_part_data(part_data_list22, "{}（{}元，{}）", name_dict=name_dict)))
    format_str_list.append(None)
    sum_data_list23, part_data_list23 = get_report_data(workbook, "挂账三年以上其他应收账款情况", [6, 8, 9],
                                                        number, [1, 6])
    sum_data_list24, part_data_list24 = get_report_data(workbook, "挂账三年以上其他应收账款情况", [],
                                                        number, [1, 8, 9])
    format_str_list.append((month, handle_num(sum_data_list23[0]), judge_up_down(sum_data_list23[1]),
                            judge_up_down(sum_data_list23[2], 1), connect_part_data(part_data_list23, name_dict=name_dict),
                            connect_part_data(part_data_list24, "{}（{}元，{}）", name_dict=name_dict)))
    sum_data_list25, part_data_list25 = get_report_data(workbook, "挂账三年以上其他应付账款情况", [6, 8, 9],
                                                        number, [1, 6])
    format_str_list.append(None)
    sum_data_list26, part_data_list26 = get_report_data(workbook, "挂账三年以上其他应付账款情况", [],
                                                        number, [1, 8, 9])
    format_str_list.append((month, handle_num(sum_data_list25[0]), judge_up_down(sum_data_list25[1]),
                            judge_up_down(sum_data_list25[2], 1), connect_part_data(part_data_list25, name_dict=name_dict),
                            connect_part_data(part_data_list26, "{}（{}元，{}）", name_dict=name_dict)))
    format_str_list.extend([None, None, None])


    sum_data_list27, part_data_list27 = get_report_data(workbook, "内部关联交易-收入确认与收货不同步", [2, 3, 5],
                                                                      number, [1, 4])
    format_str_list.append((month, sum_data_list27[2], handle_num(sum_data_list27[0]), handle_num(sum_data_list27[1]),
                            connect_part_data(part_data_list27, name_dict=name_dict)))
    format_str_list.append(None)
    sum_data_list28, part_data_list28 = get_report_data(workbook, "内部关联交易-收入确认与开票不同步", [2, 3, 4, 5],
                                                        number, [1, 2])
    sum_data_list29, part_data_list29 = get_report_data(workbook, "内部关联交易-收入确认与开票不同步", [],
                                                        number, [1, 4])
    workbook.close()
    format_str_list.append((month, handle_num(float(sum_data_list28[0]) + float(sum_data_list28[2])),
                            int(sum_data_list28[1]) + int(sum_data_list28[3]), sum_data_list28[1], handle_num(sum_data_list28[0]),
                            connect_part_data(part_data_list28, name_dict=name_dict), sum_data_list28[3], handle_num(sum_data_list28[2]),
                            connect_part_data(part_data_list29, name_dict=name_dict)))
    format_str_list.append(None)
    document = docx.Document()
    header_paragraph = document.add_paragraph()
    header_run = header_paragraph.add_run(f"南瑞集团{year}年{month}月财务监督报告")
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.line_spacing = Pt(29)
    header_paragraph.paragraph_format.space_after = Pt(0)
    p1 = document.add_paragraph()
    p1_run = p1.add_run("一、常规监督问题情况")
    p1.paragraph_format.line_spacing = Pt(29)
    p1.paragraph_format.space_after = Pt(0)
    with open("/home/cwy/static_files/report_template.txt") as f:
        data_list = f.readlines()
        # print(len(data_list), len(format_str_list))
        for data, format_str in zip(data_list, format_str_list):
            # print(data.strip())
            # print(format_str)
            # print("*"*20)
            if format_str:
                data = data.format(*format_str)
            p = document.add_paragraph()
            if data.startswith(("（", "1", "2", "3", "4", "5", "6")):
                p_run = p.add_run(data.strip())
                p_run.bold = True
                p_run.font.size = Pt(16)
                p_run.font.name = '方正楷体_GBK'
                p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正楷体_GBK')
            else:
                p_run = p.add_run(data.strip())
                p_run.font.size = Pt(16)
                p_run.font.name = '仿宋_GB2312'
                p_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = Pt(29)
            p.paragraph_format.space_after = Pt(0)
    header_run.font.name = '黑体'
    header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    header_run.font.size = Pt(18)
    header_run.bold = True
    p1_run.bold = True
    p1_run.font.name = '黑体'
    p1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p1_run.font.size = Pt(16)
    document.save(result_path)
    return True, True





